"""
MS Word(.docx) generator — 송출내역 출력 메뉴의 각 보고서를 Word로 저장.
PDF(pdf_generator.py)와 동일한 데이터를 사용해 표 중심으로 생성한다.
"""

import os
import calendar
from datetime import datetime, date as date_type

from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import (
    REPORT_MONTHLY_DIR, REPORT_DAILY_DIR, REPORT_SUMMARY_DIR, REPORT_SUBTITLE_DIR,
)

_KO_FONT = "Malgun Gothic"
_DAY_OF_WEEK_KO = ["월", "화", "수", "목", "금", "토", "일"]


# ── 공통 헬퍼 ────────────────────────────────────────────────────────────────

def _new_doc(body_font: str | None = None) -> Document:
    """한글 폰트가 적용된 새 문서. body_font 지정 시 본문 기본 폰트로 사용."""
    ko = body_font or _KO_FONT
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = ko
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), ko)
    rfonts.set(qn("w:ascii"), ko)
    rfonts.set(qn("w:hAnsi"), ko)
    # Word 기본 스타일의 문단 간격/줄간격(보통 8pt·1.08배)을 0으로 낮춘다.
    # 그대로 두면 표의 행(문단)마다 여분 간격이 누적되어, PDF에서는 한 페이지에
    # 들어가는 표(월 31행 등)가 Word에서는 여러 페이지로 넘어간다.
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    return doc


def _apply_run_font(run, font_name: str):
    """run에 지정 폰트를 라틴/한글(eastAsia) 모두 적용."""
    run.font.name = font_name
    rf = run._element.rPr.rFonts
    rf.set(qn("w:eastAsia"), font_name)
    rf.set(qn("w:ascii"), font_name)
    rf.set(qn("w:hAnsi"), font_name)


def _title(doc: Document, text: str, size: int = 16, font: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    _apply_run_font(run, font or _KO_FONT)
    return p


def _set_cell(cell, text, bold=False, align="center", size=9, font=None):
    p = cell.paragraphs[0]
    # 셀에 남아있는 run 제거 후 새로 작성 (앞쪽 빈 run 방지)
    for _r in list(p.runs):
        _r._element.getparent().remove(_r._element)
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("" if text is None else str(text))
    _apply_run_font(run, font or _KO_FONT)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    return cell


def _set_page(doc: Document, margin_mm: int = 15):
    """
    A4 페이지 + PDF와 동일한 여백으로 설정. PDF 쪽 col_w(mm) 값을 그대로 재사용할 수
    있도록 여백을 맞춰, 표 폭 비율이 PDF와 일치하게 한다.
    """
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)


def _set_col_widths(table, widths_mm: list[float]):
    """
    표의 열 폭을 PDF와 동일한 비율(mm)로 고정한다.
    Word의 고정(fixed) 레이아웃에서 실제 열 폭은 각 셀의 tcW가 아니라 표의
    tblGrid(gridCol) 값으로 결정되므로, 둘 다 갱신해야 한다. tblGrid만 빠뜨리면
    표마다 렌더링 폭이 달라지거나(표 간 폭 불일치), 의도한 폭보다 좁게 그려져
    텍스트가 줄바꿈되는 문제가 생긴다.
    """
    table.autofit = False
    table.allow_autofit = False
    widths = [Mm(w) for w in widths_mm]

    for row in table.rows:
        for idx, w in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = w

    tblGrid = table._tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        grid_cols = tblGrid.findall(qn("w:gridCol"))
        for gc, w in zip(grid_cols, widths):
            gc.set(qn("w:w"), str(int(w.twips)))


def _shade_row(row, hex_color: str):
    """표 행 전체 셀에 배경색(음영) 적용."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:shd"))
        if old is not None:
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _thick_border(row, side: str = "top", size: int = 18, color: str = "3A3A3A"):
    """행의 지정 변(top/bottom)에 굵은 테두리 적용. size 단위: 1/8 pt."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcPr.append(borders)
        old = borders.find(qn(f"w:{side}"))
        if old is not None:
            borders.remove(old)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def _thick_top_border(row, size: int = 18, color: str = "3A3A3A"):
    _thick_border(row, "top", size, color)


def _row_side_border(row, side: str, val: str = "single", size: int = 6, color: str = "000000"):
    """행의 지정 변에 테두리(스타일 지정). val: single/dotted/dashed 등."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcPr.append(borders)
        old = borders.find(qn(f"w:{side}"))
        if old is not None:
            borders.remove(old)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def _remove_table_borders(table):
    """표 전체 테두리 제거(격자 없음)."""
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def _hrule(doc: Document, size: int = 16, color: str = "000000",
           space_before: float = 0, space_after: float = 2, side: str = "bottom"):
    """가로 구분선(문단 테두리). side=bottom/top. 얇게 표시."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)
    p.add_run("").font.size = Pt(1)   # 높이 최소화
    return p


def _white_header(row, bg_hex: str = "3A3A3A"):
    """헤더 행: 진한 배경 + 흰 글자."""
    _shade_row(row, bg_hex)
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_floating_seal(paragraph, img_path, w_mm: float, h_mm: float,
                       x_mm: float, y_mm: float):
    """
    직인 이미지를 '떠있는(앞에 표시)' 이미지로 넣어 글자와 겹치게 한다.
    x_mm: 페이지 왼쪽 기준 가로 절대위치(결정적), y_mm: 문단 기준 세로 오프셋(음수=위로).
    """
    EMU = 36000  # 1mm
    run = paragraph.add_run()
    run.add_picture(img_path, width=Mm(w_mm), height=Mm(h_mm))
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                 ("simplePos", "0"), ("relativeHeight", "251658240"),
                 ("behindDoc", "0"), ("locked", "0"), ("layoutInCell", "0"),
                 ("allowOverlap", "1")):
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0")
    anchor.append(sp)
    # 가로: 페이지 기준 절대위치(셀/열 기준은 Word가 무시하는 경우가 있어 page로 고정)
    posH = OxmlElement("wp:positionH"); posH.set("relativeFrom", "page")
    oh = OxmlElement("wp:posOffset"); oh.text = str(int(x_mm * EMU)); posH.append(oh)
    anchor.append(posH)
    posV = OxmlElement("wp:positionV"); posV.set("relativeFrom", "paragraph")
    ov = OxmlElement("wp:posOffset"); ov.text = str(int(y_mm * EMU)); posV.append(ov)
    anchor.append(posV)
    anchor.append(extent)
    ee = OxmlElement("wp:effectExtent")
    for a in ("l", "t", "r", "b"):
        ee.set(a, "0")
    anchor.append(ee)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(docPr)
    anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)


def _grid_table(doc: Document, headers: list[str], font: str | None = None):
    """머리글 행이 있는 격자 표 생성 (좌측 정렬 — 여러 표의 좌측 여백을 일치시킴)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, font=font)
    return table


def _compact_table(table, row_h_mm: float = 5.0, exact: bool = True):
    """
    표를 조밀하게 만든다: 셀 상하 여백 0, 셀 안쪽 여백 최소, 행 높이 고정.
    Word 기본 행 높이가 커서 긴 표(월 31행 등)가 여러 페이지로 넘어가는 것을 막는다.
    """
    tblPr = table._tbl.tblPr
    # 셀 안쪽 여백(상/하 0, 좌/우 최소)
    old = tblPr.find(qn("w:tblCellMar"))
    if old is not None:
        tblPr.remove(old)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 0), ("bottom", 0), ("left", 40), ("right", 40)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)

    rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST
    for row in table.rows:
        row.height = Mm(row_h_mm)
        row.height_rule = rule


def _estimate_wrap_lines(text: str, col_mm: float = 68.0,
                         font_pt: float = 9.0, cell_pad_mm: float = 2.0) -> int:
    """
    Word는 렌더링 높이를 직접 잴 수 없어, 셀 폭과 글자 수로 줄바꿈 줄 수를 추정한다.
    실측(TV칸 68mm, 9pt에서 7회 이상이면 2줄)에 맞춰 글자폭을 다소 넉넉히(≈0.62em) 잡아,
    두줄 날짜를 과소 집계해 페이지가 넘치는 일이 없도록 한다.
    """
    import math
    if not text:
        return 1
    usable = max(1.0, col_mm - cell_pad_mm)          # 셀 좌우 안쪽 여백 제외
    char_w = font_pt * 0.62 / 72.0 * 25.4            # 0.62em 을 mm 로 환산(보수적)
    per_line = max(1, int(usable / char_w))
    return max(1, math.ceil(len(text) / per_line))


def _fit_font_pt(text: str, col_mm: float = 68.0, base: float = 9.0,
                 min_pt: float = 5.0, cell_pad_mm: float = 2.0) -> float:
    """
    text가 칸(col_mm) 안에 '한 줄'로 들어가도록 폰트 크기(pt)를 반환.
    기본(base)에서 들어가면 그대로, 넘치면 축소(최소 min_pt). 렌더 불가라 보수적 추정.
    → 모든 행이 한 줄이 되어 행 높이를 일정하게 유지(페이지 넘침 방지).
    """
    if not text:
        return base
    usable = max(1.0, col_mm - cell_pad_mm)
    w = len(text) * base * 0.62 / 72.0 * 25.4        # base 크기에서의 예상 폭(mm)
    if w <= usable:
        return base
    return max(min_pt, base * usable / w * 0.98)


def _small_gap(doc: Document, pt: float = 4):
    """표 사이 작은 간격 (빈 문단은 한 줄 높이를 차지해 페이지가 넘치므로 대체)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.font.size = Pt(pt)
    # 빈 문단 자체의 폰트 크기를 줄여 높이 최소화
    p.add_run("").font.size = Pt(pt)
    return p


def _footer_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── F-04 : 소재별 월 리포트 ──────────────────────────────────────────────────

def generate_monthly_docx(item_name, year, month, days, advertiser, settings,
                          start_date=None, end_date=None, client=None) -> str:
    company = (settings.get("company_name") or "광주문화방송").strip() or "광주문화방송"
    short = (settings.get("company_short") or "광주MBC").strip() or "광주MBC"
    # 제목/내용 폰트 이름(환경설정). 미지정 시 기본(Malgun Gothic).
    # Word는 폰트 '이름'만 기록하므로, 문서를 여는 PC에 해당 폰트가 설치돼 있어야 한다.
    body_font  = (settings.get("report_font_body_name")  or "").strip() or _KO_FONT
    title_font = (settings.get("report_font_title_name") or "").strip() or body_font
    doc = _new_doc(body_font)
    _set_page(doc, 12)   # 상·하·좌·우 여백 12mm(위=아래 대칭)

    # ── 로고(우측 상단, 있으면) ──
    logo_path = settings.get("logo_path", "")
    if logo_path and os.path.exists(logo_path):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_after = Pt(1)
        try:
            lp.add_run().add_picture(logo_path, width=Mm(19.2), height=Mm(4.4))   # 폭 60%·높이 40%
        except Exception:
            pass

    # ── 제목: 위·아래 같은 두께 가로줄 + 큰 제목(상하 여백 대칭) ──
    # 아래 줄은 '상단 테두리'로 만들어 빈 줄 없이 제목 바로 아래에 붙게 한다(대칭).
    _hrule(doc, size=14, color="000000", space_before=0, space_after=6, side="bottom")
    _title(doc, f"{short} 방송홍보 SB송출 현황", 28, font=title_font)  # 제목 아래 여백 6
    _hrule(doc, size=14, color="000000", space_before=0, space_after=4, side="top")

    # 대상 일자 목록 + 송출기간: start/end 지정 시 그 기간, 없으면 해당 월 전체
    from datetime import timedelta as _td
    if start_date and end_date:
        d_start = date_type.fromisoformat(start_date)
        d_end   = date_type.fromisoformat(end_date)
        date_list = [d_start + _td(days=i) for i in range((d_end - d_start).days + 1)]
    else:
        last_day = calendar.monthrange(year, month)[1]
        d_start = date_type(year, month, 1)
        d_end   = date_type(year, month, last_day)
        date_list = [date_type(year, month, dn) for dn in range(1, last_day + 1)]
    wd_s = _DAY_OF_WEEK_KO[d_start.weekday()]
    wd_e = _DAY_OF_WEEK_KO[d_end.weekday()]
    period = (f"{d_start.year}.{d_start.month}.{d_start.day}({wd_s})~"
              f"{d_end.year}.{d_end.month}.{d_end.day}({wd_e})")

    # 정보 표 (실선 격자 — 윗줄·아랫줄·세로줄 모두)
    info = doc.add_table(rows=4, cols=4)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    note = advertiser.get("note") or "송출시간은 방송사 상황에 따라 변동될 수 있음"
    client_name = (client or "").strip() or company   # 거래처명(미입력 시 회사명)
    rows = [
        ["거래처명", client_name, "사업자등록번호", "410-81-06350"],
        ["송출 내용", item_name, "대 표 이 사", settings.get("ceo_name", "")],
        ["송출 매체", advertiser.get("broadcast_medium", "TV"), "업 태·업 종", "서비스·방송"],
        ["송출 기간", period, "비  고", note],
    ]
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            cell = info.rows[r].cells[c]
            # 라벨 열(0,2) 가운데, 값 열(1,3) 좌측 + 한 칸 더 들여쓰기
            _set_cell(cell, val, bold=(c in (0, 2)),
                      align="left" if c in (1, 3) else "center", size=9, font=body_font)
            if c in (1, 3):
                cell.paragraphs[0].paragraph_format.left_indent = Mm(2)
    _set_col_widths(info, [18, 55, 24, 73])
    _compact_table(info, row_h_mm=6.0, exact=False)   # 비고 줄바꿈 대비 AT_LEAST
    # 셀 글자 세로 가운데 정렬(윗줄에 붙지 않게)
    for _r in info.rows:
        for _c in _r.cells:
            _c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _small_gap(doc)

    # 데이터 표 (비고 열 포함)
    table = _grid_table(doc, ["일 시", "요일", "횟수", "T V", "RADIO-AM", "RADIO-FM", "비고"], font=body_font)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER   # 정보표와 함께 가운데 정렬(좌우 여백 동일)
    _white_header(table.rows[0], "3A3A3A")        # 첫 행: 진회색 배경 + 흰 글자
    days_map = {d["date"]: d for d in days}
    total = 0
    extra_lines = 0   # TV 시간 줄바꿈으로 늘어나는 총 추가 줄 수(행 높이 계산용)
    for d in date_list:
        ds = d.strftime("%Y-%m-%d")
        wd = _DAY_OF_WEEK_KO[d.weekday()]
        info_d = days_map.get(ds)
        if info_d:
            cnt = info_d["count"]
            times = " ".join(_hhmm(t) for t in info_d["times"])
            total += cnt
        else:
            cnt, times = 0, ""
        # TV 시간이 많으면 줄바꿈하지 않고 폰트를 줄여 '한 줄'로 → 행 높이 일정
        tv_size = _fit_font_pt(times)
        cells = table.add_row().cells
        vals = [f"{d.year % 100}. {d.month}. {d.day}", wd, str(cnt) if cnt else "", times, "-", "-", ""]
        for i, v in enumerate(vals):
            # 모든 칸 가운데 정렬 + 세로 가운데. TV(i==3)만 자동 축소 폰트.
            _set_cell(cells[i], v, align="center", size=(tv_size if i == 3 else 9), font=body_font)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 합계: 2행 — (A) 매체별 합계(TV/RADIO), (B) 아래 총 합계
    row_a = table.add_row()
    row_b = table.add_row()
    ra, rb = row_a.cells, row_b.cells
    _set_cell(ra[0], "총 계", bold=True, size=9, font=body_font)
    _set_cell(ra[1], f"{len(date_list)} 일", bold=True, size=9, font=body_font)
    _set_cell(ra[2], f"{total} 회", bold=True, size=9, font=body_font)   # 횟수+TV 합계
    _set_cell(ra[4], "-", bold=True, size=9, font=body_font)
    _set_cell(ra[5], "-", bold=True, size=9, font=body_font)
    _set_cell(rb[2], f"총 {total} 회", bold=True, size=9, font=body_font)  # 총 합계
    _set_col_widths(table, [22, 10, 12, 68, 22, 22, 14])
    # 병합: 횟수+TV(윗줄), 총합계(아랫줄 넓게), '총 계'·일수 세로 병합
    ra[2].merge(ra[3])
    m = rb[2]
    for _i in (3, 4, 5, 6):
        m = m.merge(rb[_i])
    ra[0].merge(rb[0])
    ra[1].merge(rb[1])
    _shade_row(row_a, "EAE3D0")
    _shade_row(row_b, "EAE3D0")
    # 모든 셀 세로 가운데 정렬(글자가 위/아래 선에 붙지 않게)
    for _row in table.rows:
        for _cell in _row.cells:
            _cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── 행 높이 자동 계산 (PDF처럼 위/아래 여백 균형 + 페이지 채우기) ──
    # 가용 세로 높이에서 제목·정보표·간격·푸터(추정 고정분)를 뺀 나머지를
    # 데이터 표의 각 행에 고르게 나눠, 내용이 적어도 표가 페이지를 채우게 한다.
    # 줄바꿈으로 늘어나는 행(extra_lines)만큼 높이를 먼저 확보해 넘침을 막는다.
    num_rows = len(date_list) + 3           # 머리글 + 일자 + 총계 2줄
    usable_h = 297.0 - 24.0                 # A4 높이 - 상하 여백(12+12, 대칭)
    # 제목+정보표+간격+푸터의 실제 렌더 높이는 서버에서 정확히 잴 수 없어,
    # 행 높이 6.6mm가 2페이지로 넘친(=푸터 밀림) 실측을 근거로 오버헤드를 넉넉히
    # 잡아 한 페이지를 보장한다. (row_h ≈ 5.9mm 수준으로 수렴)
    # 상단 고정분(로고·제목·가로줄·정보표·간격·푸터·꼬리문단·푸터앞 빈줄). 넉넉히 잡아
    # 표를 짧게 → 합계·푸터가 항상 1페이지에 들어오게 한다. (아래 여백은 다소 남을 수 있음)
    overhead_mm = 98.0
    line_h_mm = 3.8                         # 9pt 한 줄 높이(줄바꿈 1줄당 추가 높이)
    # 두줄 날짜(extra_lines)만큼 높이를 먼저 확보하고 나머지를 전 행에 고르게 분배.
    #  → 줄바꿈 많으면 행높이 자동 축소(1페이지), 없으면 넉넉히 채워 위·아래 여백 균형.
    avail_table = usable_h - overhead_mm - extra_lines * line_h_mm
    row_h = avail_table / num_rows
    row_h = max(3.8, min(6.5, row_h))       # 최소=조밀(넘침 방지), 최대=과도한 벌어짐 방지
    # AT_LEAST: 계산 높이로 채우되, 줄바꿈 행은 잘리지 않고 더 늘어남
    _compact_table(table, row_h_mm=row_h, exact=False)

    # 푸터 앞 간격(표와 한 줄 더 띄움 — PDF와 비슷하게)
    _small_gap(doc)
    doc.add_paragraph().add_run("").font.size = Pt(11)   # 빈 한 줄

    # 푸터: 확인 문구 + 회사명(표 우측 끝보다 안쪽) + 회사명 위 직인 겹침
    # 3열 [확인문구 | 회사명 | 우측 여백] → 회사명이 표 우측 끝보다 안쪽에 놓임
    ftbl = doc.add_table(rows=1, cols=3)
    ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc, rc, _pad = ftbl.rows[0].cells
    _confirm = f"위와 같이 방송 송출 완료하였음을 확인함 ({d_end.year}. {d_end.month}. {d_end.day}.)"
    _set_cell(lc, _confirm, align="center", size=10.5, font=body_font)
    _set_cell(rc, f"{company}(주)", bold=True, align="right", size=11, font=body_font)
    _set_cell(_pad, "", size=11, font=body_font)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 확인 문구가 한 줄에 들어가도록 좌측 칸을 넓게, 회사명 칸은 우측 끝에서 안쪽
    _set_col_widths(ftbl, [128, 38, 4])   # 합계 170mm

    # 직인: 회사명 칸 문단에 '떠있는' 이미지로 '(주)' 글자 끝에 겹치게 한다.
    # x_mm: 좌우(작을수록 왼쪽), y_mm: 상하(음수=위, 클수록 아래).
    seal_path = settings.get("seal_path", "")
    if seal_path and os.path.exists(seal_path):
        try:
            # x_mm: 페이지 왼쪽 기준 가로(작을수록 왼쪽), y_mm: 세로(작을수록 위로)
            # 글자를 가리지 않도록 오른쪽으로: '(주)' 끝에 살짝만 겹치게
            _add_floating_seal(rc.paragraphs[0], seal_path, w_mm=15, h_mm=15,
                               x_mm=184, y_mm=-5)
        except Exception:
            pass

    safe_item = str(item_name).replace("/", "_").replace("\\", "_")[:80]
    out = str(REPORT_MONTHLY_DIR / f"SB송출현황_{safe_item}_{year}{month:02d}.docx")
    doc.save(out)
    return out


def _hhmm(time_str: str) -> str:
    """'HH:MM:SS' → 4자리 HHMM (실제 시계 시각 그대로, 화면과 동일하게 +24 변환 없음)."""
    p = time_str.split(":")
    if len(p) < 2:
        return time_str
    h, m = int(p[0]), int(p[1])
    return f"{h:02d}{m:02d}"


# ── F-06 : 방송 운행표 ───────────────────────────────────────────────────────

def generate_daily_docx(date, items, settings) -> str:
    doc = _new_doc()
    _set_page(doc, 15)
    _title(doc, f"방송 운행표  —  {date}", 14)

    if not items:
        doc.add_paragraph("해당 날짜의 SB 송출 내역이 없습니다.")
    else:
        _SKY   = "DCE6F1"   # 프로그램 행 음영
        _GREEN = "E2EFDA"   # 광고/광고그룹 행 음영

        def _is_end_notice(prog):
            return "방송 종료" in (prog or "") or "방송종료" in (prog or "")

        table = _grid_table(doc, ["방송시작시간", "프로그램명", "소재종류", "SB 소재 제목"])
        for i, r in enumerate(items):
            lbl = r.get("content_type_label", "")
            prog = r.get("program_block", "")
            is_ct = lbl == "이어서"   # '이어서' 행 굵게
            row = table.add_row()
            cells = row.cells
            _set_cell(cells[0], r.get("broadcast_time_display") or r["broadcast_time"], bold=is_ct)
            _set_cell(cells[1], prog, bold=is_ct)
            _set_cell(cells[2], lbl, bold=is_ct)
            _set_cell(cells[3], r.get("item_name_raw", ""), bold=is_ct, align="left")
            if is_ct:
                _thick_top_border(row)          # 프로그램 그룹 시작 → 위쪽 굵은 구분선
            elif lbl == "프로그램" and "방송순서" not in (prog or ""):
                # 첫부분 '방송순서 안내'는 프로그램이어도 하늘색 음영 제외
                _shade_row(row, _SKY)
            elif lbl in ("광고", "광고그룹"):
                _shade_row(row, _GREEN)
            # 방송종료 안내 시작 행 → 위쪽 굵은 구분선
            if not is_ct and _is_end_notice(prog) \
                    and not (i > 0 and _is_end_notice(items[i - 1].get("program_block", ""))):
                _thick_top_border(row)
        _set_col_widths(table, [28, 48, 22, 82])

        # 표 외곽선 굵게: 맨 윗줄(헤더 위) / 헤더(제목) 아래 / 맨 밑줄
        _thick_border(table.rows[0], "top")
        _thick_border(table.rows[0], "bottom")
        _thick_border(table.rows[-1], "bottom")

    _footer_note(doc, f"총 {len(items)}건 | 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    out = str(REPORT_DAILY_DIR / f"방송운행표_{date.replace('-', '')}.docx")
    doc.save(out)
    return out


# ── 일일 운행표 / 일일 ID 운행표 ──────────────────────────────────────────────

def generate_daily_summary_docx(date, type_label, items) -> str:
    doc = _new_doc()
    _set_page(doc, 15)
    title = "일일 ID 운행표" if type_label == "ID" else "일일 운행표"
    _title(doc, f"{title}  —  {date}", 14)

    if not items:
        doc.add_paragraph("해당 날짜의 송출 내역이 없습니다.")
    else:
        table = _grid_table(doc, ["소재명", "총횟수", "SA", "A", "B", "C"])
        for r in items:
            cells = table.add_row().cells
            _set_cell(cells[0], r["item_name"], align="left")
            _set_cell(cells[1], f"{r['total_count']}회")
            _set_cell(cells[2], r.get("sa") or 0)
            _set_cell(cells[3], r.get("a") or 0)
            _set_cell(cells[4], r.get("b") or 0)
            _set_cell(cells[5], r.get("c") or 0)
        _set_col_widths(table, [70, 22, 22, 22, 22, 22])

    total = sum(r["total_count"] for r in items)
    _footer_note(doc, f"총 {total}회 | 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    prefix = "일일ID운행표" if type_label == "ID" else "일일운행표"
    out = str(REPORT_SUMMARY_DIR / f"{prefix}_{date.replace('-', '')}.docx")
    doc.save(out)
    return out


# ── 흘림자막·공익·재난 송출내역 ────────────────────────────────────────────────

def _fmt_dur(sec) -> str:
    if not sec:
        return ""
    sec = int(sec)
    if sec < 60:
        return f'{sec}"'
    return f"{sec // 60}'{sec % 60:02d}\""


def _hhmm_ko(t: str) -> str:
    if not t:
        return ""
    p = t.split(":")
    try:
        return f"{int(p[0]):02d}시{int(p[1]):02d}분"
    except (ValueError, IndexError):
        return t


def generate_subtitle_campaign_docx(data) -> str:
    doc = _new_doc()
    _set_page(doc, 12)
    date = data["date"]
    _title(doc, "□ 흘림자막 및 공익광고/재난피해 사전예방 송출내역 □", 15)
    try:
        d = datetime.strptime(date, "%Y-%m-%d")
        wd = _DAY_OF_WEEK_KO[d.weekday()]
        _title(doc, f"{d.year}년 {d.month}월 {d.day}일 ({wd}) (00:00 ~ 24:00)", 12)
    except ValueError:
        _title(doc, f"{date} (00:00 ~ 24:00)", 12)

    def _section(name, headers, rows, widths_mm):
        doc.add_paragraph().add_run(f"□ {name}").bold = True
        t = _grid_table(doc, headers)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                _set_cell(cells[i], v, align="left" if i == 1 else "center")
        _set_col_widths(t, widths_mm)
        doc.add_paragraph()

    # 1. UHD방송홍보 (영상/자막)
    uhd_v, uhd_s = data["uhd_video"], data["uhd_sub"]
    rows = []
    for i in range(max(len(uhd_v), len(uhd_s), 1)):
        v = uhd_v[i] if i < len(uhd_v) else {}
        s = uhd_s[i] if i < len(uhd_s) else {}
        rows.append([_hhmm_ko(v.get("time", "")), v.get("program", ""),
                     _hhmm_ko(s.get("time", "")), s.get("program", "")])
    _section("UHD방송홍보", ["송출시간(영상)", "프로그램", "송출시간(자막)", "프로그램"], rows,
              [30, 63, 30, 63])

    # 2. TV직접수신 (2쌍)
    tv = data["tv_direct"]
    rows = []
    for i in range(max((len(tv) + 1) // 2, 1)):
        a = tv[2 * i] if 2 * i < len(tv) else {}
        b = tv[2 * i + 1] if 2 * i + 1 < len(tv) else {}
        rows.append([_hhmm_ko(a.get("time", "")), a.get("program", ""),
                     _hhmm_ko(b.get("time", "")), b.get("program", "")])
    _section("TV직접수신", ["송출시간(자막)", "프로그램", "송출시간(자막)", "프로그램"], rows,
              [30, 63, 30, 63])

    # 3. 시청자의견
    vo = data["viewer_opinion"]
    rows = [[_hhmm_ko(r.get("time", "")), r.get("program", "")]
            for r in (vo or [{}])]
    _section("시청자의견 (주1회 목요일)", ["송출시간(자막)", "프로그램"], rows,
              [35, 151])

    # 4. 공익광고
    worker = data.get("campaign_worker", "")
    camp = data["campaign"]
    rows = [[r.get("time", ""), r.get("program", ""), _fmt_dur(r.get("duration")),
             r.get("grade", ""), worker] for r in camp]
    if not rows:
        rows = [["", "", "", "", ""]]
    _section("공익광고 송출내역 (본사 포함)",
             ["방송시간", "프로그램", "초수", "시급", "근무자"], rows,
              [28, 94, 22, 18, 24])

    # 5. 재난피해
    dis = data["disaster"]
    rows = [[r.get("time", ""), r.get("program", ""), _fmt_dur(r.get("duration")), worker]
            for r in dis]
    if not rows:
        rows = [["", "", "", ""]]
    _section("재난피해 사전예방 프로그램 송출내역 (본사 포함)",
             ["방송시간", "프로그램", "초수", "근무자"], rows,
              [28, 112, 22, 24])

    _footer_note(doc, f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    out = str(REPORT_SUBTITLE_DIR / f"흘림자막공익재난송출내역_{date.replace('-', '')}.docx")
    doc.save(out)
    return out
